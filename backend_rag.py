# backend_rag.py
"""
Combined backend for RAG per-user-per-chat (updated).
- Per-user-per-chat vector stores (vectors/{user_id}/{thread_id}/...)
- Ingest, retrieval, and DB mappings
- quick_analyze and get_term_context utilities
- Semantic-aware chunking (heading/paragraph aware) with sliding-window fallback

Notes:
- OCR for images and scanned PDFs uses Google Vision.
- For PDFs:
  - Preferred: async batch OCR via Google Vision using GCS (VISION_GCS_BUCKET env var).
  - Fallback: local rasterization (pdf2image) + per-page Vision calls (requires poppler).
- GOOGLE_APPLICATION_CREDENTIALS is read from env and used to construct Vision/Storage clients.
"""

import os
import json
import uuid
import sqlite3
import shutil
import math
import re
import subprocess
import time
from pathlib import Path
from typing import Optional, List, Dict, Any, TypedDict, Annotated, Tuple
from datetime import datetime
from collections import Counter
from io import BytesIO

# load env
from dotenv import load_dotenv
load_dotenv()

# --- Optional dependencies (fail fast with helpful message only on import time for critical libs) ---
try:
    from langgraph.graph import StateGraph, START, END
    from langgraph.checkpoint.sqlite import SqliteSaver
    from langgraph.graph.message import add_messages
    from langchain_core.messages import BaseMessage, HumanMessage, SystemMessage
    from langchain_google_genai import ChatGoogleGenerativeAI
except Exception as e:
    raise ImportError(f"Install langgraph/langchain-google-genai packages. Error: {e}")

# RAG libs (required)
try:
    from sentence_transformers import SentenceTransformer
    try:
        from sentence_transformers import CrossEncoder
    except Exception:
        CrossEncoder = None
    import numpy as np
    import faiss
    import PyPDF2
except Exception as e:
    raise ImportError(f"Install dependencies: sentence-transformers, faiss-cpu, PyPDF2, numpy. Error: {e}")

# optional docx
try:
    import docx
except Exception:
    docx = None

# image lib PIL
try:
    from PIL import Image
except Exception:
    Image = None

# pdf2image optional (local rasterization fallback)
try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

# Google Vision + Storage (lazy clients created later)
try:
    from google.cloud import vision_v1 as vision
    from google.cloud import storage as gcs
    from google.oauth2 import service_account
    _vision_import_error = None
except Exception as e:
    vision = None
    gcs = None
    service_account = None
    _vision_import_error = e

# optional wordfreq
try:
    from wordfreq import zipf_frequency
except Exception:
    zipf_frequency = None

# ---------- Paths & DB ----------
BASE_DIR = Path(__file__).parent
VECTOR_DIR = BASE_DIR / "vectors"
VECTOR_DIR.mkdir(exist_ok=True, parents=True)
DB_PATH = BASE_DIR / "chatbot.db"

conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
checkpointer = SqliteSaver(conn=conn)

# ---------- Prompt templates (restored to your specified long prompt) ----------
STRICT_SYSTEM_PROMPT_TEMPLATE = (
    "You are a meticulous legal-document assistant. Use ONLY the provided document excerpts to answer. "
    "Do NOT use external knowledge or make up facts unless explicitly asked; when you do use general legal principles, label them clearly as GENERAL LEGAL INFORMATION and separate from document-based findings.\n\n"

    "Answer the user's question by producing the following sections in this exact order and format:\n\n"

    "PLAIN ANSWER:\n"
    "  - 1–3 short sentences in plain English aimed at a non-lawyer. If the question is yes/no, start with 'Yes' or 'No' and then add 1 sentence of brief explanation.\n\n"

    "FACTS (up to 3):\n"
    "  - For each fact, provide a one-line factual statement DIRECTLY SUPPORTED by the excerpts, followed by a short SOURCE tag in parentheses (File/ChunkID/Page) and a quoted snippet (<=200 characters). If you cannot find any facts, write: FACTS: Not stated in document.\n\n"

    "EVIDENCE / EXCERPTS:\n"
    "  - List the specific excerpt(s) used (at most 2). For each: source id or file name; a short snippet in quotes (<=250 chars); and chunk id or page if available. Do not paste more than 250 characters per excerpt.\n\n"

    "ASSESSMENT:\n"
    "  - CONFIDENCE: High / Medium / Low. Provide one short reason for the confidence (e.g., 'direct quote present', 'requires inference from two clauses', 'conflicting language in excerpts'). If you made any inference, prepend 'INFERENCE:' and explain the premises and the specific excerpts used (1–2 sentences).\n\n"

    "CONFLICTS:\n"
    "  - If excerpts contain contradictory statements, list the conflicting lines and cite sources. If none, write: CONFLICTS: None.\n\n"

    "NEXT STEPS:\n"
    "  - Give 1–2 short actionable next steps (e.g., 'Consult counsel', 'Upload full contract', 'Check clause X on page Y'). If the answer could affect legal rights or obligations, include a strong recommendation to obtain lawyer review.\n\n"

    "IMPORTANT RULES — ALWAYS FOLLOW:\n"
    "- Use ONLY the provided excerpts. If the answer is not determinable from the excerpts, say 'Not stated in document' and DO NOT guess.\n"
    "- If you must infer, prepend 'INFERENCE:' and explain reasoning and show the excerpts used.\n"
    "- Do NOT provide boilerplate legal advice; when appropriate, recommend lawyer review.\n"
    "- Keep language plain and concise.\n\n"

    "Document context:\n{context}\n\nNow answer the user's question succinctly and accurately following the specified format."
)

FALLBACK_SYSTEM_PROMPT = (
    "No document context provided. You are a legal-document assistant.\n\n"
    "If the user question requires the document to answer, reply: 'Cannot determine from available information — please upload the document or provide the clause.'\n\n"
    "If the user asks for general legal information (not document-specific), provide a short GENERAL LEGAL INFORMATION section (1-2 sentences), label it clearly, and include: 'This is general information, not legal advice.'\n\n"
    "Then give NEXT STEPS: one or two practical actions (e.g., 'Upload the contract', 'Ask for clause X', 'Consult a lawyer')."
)

def _build_strict_system_prompt(context: str, max_context_chars: int = 5000) -> str:
    if not context:
        context = "(no excerpts provided)"
    if len(context) > max_context_chars:
        head = context[:max_context_chars//2]
        tail = context[-max_context_chars//2:]
        context = head + "\n\n...[TRUNCATED]...\n\n" + tail
    return STRICT_SYSTEM_PROMPT_TEMPLATE.format(context=context)

# ---------- DB schema helpers ----------
def _backup_db():
    try:
        bak = DB_PATH.with_suffix(".db.bak")
        if DB_PATH.exists() and not bak.exists():
            shutil.copy2(DB_PATH, bak)
    except Exception:
        pass

def _table_columns(table_name: str):
    cur = conn.cursor()
    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
    if not cur.fetchone():
        return None
    cur.execute(f"PRAGMA table_info({table_name})")
    rows = cur.fetchall()
    return [r[1] for r in rows]

def _ensure_schema():
    if DB_PATH.exists():
        _backup_db()
    cur = conn.cursor()
    users_cols = _table_columns("users")
    if users_cols is None:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                user_id TEXT PRIMARY KEY,
                username TEXT,
                created_at TEXT
            )
        """)
        conn.commit()
    threads_cols = _table_columns("threads")
    if threads_cols is None:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS threads (
                thread_id TEXT PRIMARY KEY,
                user_id TEXT,
                file_name TEXT,
                filepath TEXT,
                ingested_at TEXT
            )
        """)
        conn.commit()
    conn.commit()

_try_schema_err = None
try:
    _ensure_schema()
except Exception as _e:
    _try_schema_err = _e
    print("Warning: _ensure_schema failed:", _e)

# ---------- Model ----------
DEFAULT_GOOGLE_MODEL = os.getenv('GOOGLE_MODEL', 'gemini-2.5-flash')
DEFAULT_MODEL_TEMPERATURE = float(os.getenv('MODEL_TEMPERATURE', 0.0))
model = ChatGoogleGenerativeAI(
    model=DEFAULT_GOOGLE_MODEL,
    temperature=DEFAULT_MODEL_TEMPERATURE,
    google_api_key=os.getenv('GOOGLE_API_KEY')
)

# ---------- Embedding model ----------
EMBEDDING_MODEL_NAME = os.getenv('EMBEDDING_MODEL', "all-MiniLM-L6-v2")
_embed_model = SentenceTransformer(EMBEDDING_MODEL_NAME)

# ---------- Hybrid retrieval / reranker config ----------
ANN_TOP_K = int(os.getenv("ANN_TOP_K", "100"))
FINAL_TOP_K = int(os.getenv("FINAL_TOP_K", "5"))
CROSS_ENCODER_MODEL = os.getenv("CROSS_ENCODER_MODEL", "")

_cross_encoder = None
if CrossEncoder is not None and CROSS_ENCODER_MODEL:
    try:
        _cross_encoder = CrossEncoder(CROSS_ENCODER_MODEL)
    except Exception:
        _cross_encoder = None

RERANKER_AVAILABLE = _cross_encoder is not None

# ---------- Vision + GCS clients (lazy) ----------
_vision_client = None
_gcs_client = None
VISION_AVAILABLE = vision is not None and _vision_import_error is None

def _get_vision_and_storage_clients():
    """
    Create vision and storage clients using GOOGLE_APPLICATION_CREDENTIALS if present.
    Returns (vision_client, storage_client) or (None, None) if unavailable.
    """
    global _vision_client, _gcs_client
    if not VISION_AVAILABLE:
        return None, None
    if _vision_client is not None and _gcs_client is not None:
        return _vision_client, _gcs_client
    key_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
    creds = None
    try:
        if key_path and os.path.exists(key_path) and service_account is not None:
            creds = service_account.Credentials.from_service_account_file(key_path)
            _vision_client = vision.ImageAnnotatorClient(credentials=creds)
            if gcs is not None:
                # create storage client with same creds
                _gcs_client = gcs.Client(credentials=creds, project=creds.project_id)
            else:
                _gcs_client = None
        else:
            # rely on default ADC (if running on GCE/GKE or gcloud auth)
            _vision_client = vision.ImageAnnotatorClient()
            _gcs_client = gcs.Client() if gcs is not None else None
        return _vision_client, _gcs_client
    except Exception:
        _vision_client = None
        _gcs_client = None
        return None, None

# ---------- Per-user-per-chat path helpers ----------
def _chat_paths(user_id: Optional[str], thread_id: str):
    uid = str(user_id) if user_id else "anonymous"
    chat_dir = VECTOR_DIR / uid / str(thread_id)
    chat_dir.mkdir(parents=True, exist_ok=True)
    faiss_path = chat_dir / "faiss.index"
    meta_path = chat_dir / "metadata.json"
    file_record = chat_dir / "ingested_file.json"
    return faiss_path, meta_path, file_record

def _ensure_index_for_chat(user_id: Optional[str], thread_id: str):
    faiss_path, meta_path, _ = _chat_paths(user_id, thread_id)
    if faiss_path.exists() and meta_path.exists():
        try:
            index = faiss.read_index(str(faiss_path))
            with open(meta_path, "r", encoding="utf-8") as f:
                metadata = json.load(f)
            return index, metadata
        except Exception:
            pass
    d = _embed_model.get_sentence_embedding_dimension()
    index = faiss.IndexFlatL2(d)
    metadata = {}
    faiss.write_index(index, str(faiss_path))
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)
    return index, metadata

def _persist_index_for_chat(user_id: Optional[str], thread_id: str, index, metadata: dict):
    faiss_path, meta_path, _ = _chat_paths(user_id, thread_id)
    faiss.write_index(index, str(faiss_path))
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

# ---------- OCR helpers using Google Vision (GCS async for PDFs) ----------
def ocr_image_with_google_vision_bytes(content: bytes) -> str:
    vclient, _ = _get_vision_and_storage_clients()
    if vclient is None:
        return ""
    try:
        image = vision.Image(content=content)
        resp = vclient.document_text_detection(image=image)
        if resp.error and getattr(resp.error, "message", None):
            return ""
        fta = getattr(resp, "full_text_annotation", None)
        if fta and getattr(fta, "text", None):
            return fta.text
        # fallback to text_annotations
        if resp.text_annotations and len(resp.text_annotations) > 0:
            return resp.text_annotations[0].description or ""
        return ""
    except Exception:
        return ""

def ocr_image_with_google_vision(image_path: str) -> str:
    try:
        with open(image_path, "rb") as f:
            return ocr_image_with_google_vision_bytes(f.read())
    except Exception:
        return ""

def _upload_file_to_gcs(bucket_name: str, source_path: str, dest_blob_name: str) -> Optional[str]:
    """Upload local file to GCS. Returns gs:// URI or None."""
    _, storage_client = _get_vision_and_storage_clients()
    if storage_client is None:
        return None
    try:
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(dest_blob_name)
        blob.upload_from_filename(source_path)
        return f"gs://{bucket_name}/{dest_blob_name}"
    except Exception:
        return None

def _read_gcs_json_outputs(bucket_name: str, prefix: str) -> List[dict]:
    """List and read JSON output files under prefix (folder) and return parsed JSON objects."""
    _, storage_client = _get_vision_and_storage_clients()
    if storage_client is None:
        return []
    out = []
    try:
        bucket = storage_client.bucket(bucket_name)
        blobs = list(storage_client.list_blobs(bucket_name, prefix=prefix))
        for b in blobs:
            if not b.name.endswith(".json"):
                continue
            data = b.download_as_bytes()
            try:
                parsed = json.loads(data.decode("utf-8"))
                out.append(parsed)
            except Exception:
                try:
                    parsed = json.loads(data)
                    out.append(parsed)
                except Exception:
                    continue
        return out
    except Exception:
        return []

def ocr_pdf_with_google_vision_async_gcs(pdf_path: str, max_pages: int = 100) -> str:
    """
    Use Google Vision async_batch_annotate_files by uploading PDF to GCS and reading JSON output.
    Returns aggregated text or empty string on failure.
    Requires VISION_GCS_BUCKET env var and service account with storage permissions.
    """
    bucket = os.environ.get("VISION_GCS_BUCKET")
    if not bucket:
        return ""
    # ensure clients
    vclient, storage_client = _get_vision_and_storage_clients()
    if vclient is None or storage_client is None:
        return ""
    try:
        # generate unique names
        uid = str(uuid.uuid4())
        gcs_src_name = f"vision_inputs/{uid}/{Path(pdf_path).name}"
        gcs_dst_prefix = f"vision_outputs/{uid}/"
        src_gs_uri = _upload_file_to_gcs(bucket, pdf_path, gcs_src_name)
        if not src_gs_uri:
            return ""

        # build request for async_batch_annotate_files
        gcs_source = vision.GcsSource(uri=src_gs_uri)
        input_config = vision.InputConfig(gcs_source=gcs_source, mime_type="application/pdf")
        gcs_destination = vision.GcsDestination(uri=f"gs://{bucket}/{gcs_dst_prefix}")
        output_config = vision.OutputConfig(gcs_destination=gcs_destination, batch_size=2)

        feature = vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)

        async_req = vision.AsyncAnnotateFileRequest(
            input_config=input_config,
            features=[feature],
            output_config=output_config
        )

        operation = vclient.async_batch_annotate_files(requests=[async_req])

        # Wait for operation to complete (with timeout)
        timeout_seconds = int(os.getenv("VISION_ASYNC_TIMEOUT", "180"))
        poll_interval = 2
        waited = 0
        while not operation.done():
            time.sleep(poll_interval)
            waited += poll_interval
            if waited > timeout_seconds:
                break

        # After operation completes, the outputs are saved to GCS prefix as JSON.
        outputs = _read_gcs_json_outputs(bucket, gcs_dst_prefix)
        text_parts = []
        for out_json in outputs:
            try:
                responses = out_json.get("responses", [])
                for r in responses:
                    fta = r.get("fullTextAnnotation")
                    if fta and isinstance(fta, dict):
                        t = fta.get("text", "")
                        if t:
                            text_parts.append(t)
                    elif "textAnnotations" in r and r["textAnnotations"]:
                        t = r["textAnnotations"][0].get("description", "")
                        if t:
                            text_parts.append(t)
            except Exception:
                continue
        return "\n\n".join(text_parts).strip()
    except Exception:
        return ""

def ocr_pdf_with_google_vision_local_pages(pdf_path: str, max_pages: int = 5) -> str:
    """
    Fallback: rasterize first max_pages via pdf2image and call Vision per-page.
    Requires pdf2image + poppler installed.
    """
    if convert_from_path is None:
        return ""
    try:
        pages = convert_from_path(pdf_path, dpi=300, first_page=1, last_page=max_pages)
        texts = []
        for p in pages:
            buf = BytesIO()
            p.save(buf, format="JPEG")
            content = buf.getvalue()
            t = ocr_image_with_google_vision_bytes(content)
            if t:
                texts.append(t)
        return "\n\n".join(texts)
    except Exception:
        return ""

# ---------- File extraction with diagnostics ----------
def _extract_text_from_pdf(path: str) -> str:
    # Try PyPDF2 text extraction (text layer)
    text = []
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                except Exception:
                    page_text = ""
                if page_text:
                    text.append(page_text)
    except Exception:
        return ""
    return "\n".join(text)

def _convert_doc_to_docx_with_soffice(doc_path: str, timeout: int = 60) -> Optional[str]:
    try:
        soffice = shutil.which("soffice")
        if soffice is None:
            possible = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
            for p in possible:
                if Path(p).exists():
                    soffice = p
                    break
        if soffice is None:
            return None
        doc_path = str(Path(doc_path).absolute())
        out_dir = str(Path(doc_path).parent.absolute())
        cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", out_dir, doc_path]
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
        if proc.returncode != 0:
            return None
        new_path = str(Path(out_dir) / (Path(doc_path).stem + ".docx"))
        return new_path if Path(new_path).exists() else None
    except Exception:
        return None

def extract_text_with_diagnostics(filepath: str, pdf_ocr_pages: int = 5) -> Dict[str, Any]:
    """
    Attempts to extract text using multiple methods and returns diagnostics.
    Returns:
      {
        "text": str,
        "source": str or None,
        "diagnostics": { step_name: {"ok": bool, "len": int, "note": str} }
      }
    """
    p = Path(filepath)
    suf = p.suffix.lower()
    diag: Dict[str, Any] = {}
    # 1) PDF: try text layer first -> preferred GCS async OCR -> local rasterize fallback
    if suf == ".pdf":
        t = _extract_text_from_pdf(filepath)
        diag["pdf_pytext"] = {"ok": bool(t and t.strip()), "len": len(t) if t else 0}
        if t and t.strip():
            return {"text": t, "source": "pdf_pytext", "diagnostics": diag}

        # If user provided VISION_GCS_BUCKET, try async GCS method
        gcs_bucket = os.environ.get("VISION_GCS_BUCKET")
        if gcs_bucket:
            diag["pdf_vision_mode"] = {"mode": "gcs_async", "bucket": gcs_bucket}
            t2 = ocr_pdf_with_google_vision_async_gcs(filepath, max_pages=pdf_ocr_pages)
            diag["pdf_vision"] = {"ok": bool(t2 and t2.strip()), "len": len(t2) if t2 else 0}
            if t2 and t2.strip():
                return {"text": t2, "source": "pdf_vision_gcs_async", "diagnostics": diag}
            # continue to local fallback
        else:
            diag["pdf_vision_mode"] = {"mode": "local_pages_fallback", "note": "VISION_GCS_BUCKET not set"}

        # local rasterization (pdf2image) fallback
        t3 = ocr_pdf_with_google_vision_local_pages(filepath, max_pages=pdf_ocr_pages)
        diag["pdf_vision_local_pages"] = {"ok": bool(t3 and t3.strip()), "len": len(t3) if t3 else 0}
        if t3 and t3.strip():
            return {"text": t3, "source": "pdf_vision_local_pages", "diagnostics": diag}

        return {"text": "", "source": None, "diagnostics": diag}

    # 2) text / md
    if suf in {".txt", ".md"}:
        try:
            t = p.read_text(encoding="utf-8", errors="ignore")
            diag["txt_read"] = {"ok": bool(t and t.strip()), "len": len(t) if t else 0}
            if t and t.strip():
                return {"text": t, "source": "txt_read", "diagnostics": diag}
        except Exception as e:
            diag["txt_read"] = {"ok": False, "len": 0, "note": str(e)}
        return {"text": "", "source": None, "diagnostics": diag}

    # 3) docx
    if suf == ".docx":
        if docx is None:
            diag["docx"] = {"ok": False, "len": 0, "note": "python-docx not installed"}
            return {"text": "", "source": None, "diagnostics": diag}
        try:
            d = docx.Document(filepath)
            paragraphs = [pp.text for pp in d.paragraphs]
            t = "\n".join(paragraphs)
            diag["docx"] = {"ok": bool(t and t.strip()), "len": len(t) if t else 0}
            if t and t.strip():
                return {"text": t, "source": "docx", "diagnostics": diag}
        except Exception as e:
            diag["docx"] = {"ok": False, "len": 0, "note": str(e)}
        return {"text": "", "source": None, "diagnostics": diag}

    # 4) .doc -> try conversion to docx
    if suf == ".doc":
        converted = _convert_doc_to_docx_with_soffice(filepath)
        diag["doc_to_docx"] = {"ok": bool(converted), "note": f"converted_path={converted}"}
        if converted:
            if docx is not None:
                try:
                    d = docx.Document(converted)
                    paragraphs = [pp.text for pp in d.paragraphs]
                    t = "\n".join(paragraphs)
                    diag["docx_after_convert"] = {"ok": bool(t and t.strip()), "len": len(t) if t else 0}
                    if t and t.strip():
                        return {"text": t, "source": "docx_converted", "diagnostics": diag}
                except Exception as e:
                    diag["docx_after_convert"] = {"ok": False, "len": 0, "note": str(e)}
            try:
                raw = Path(converted).read_text(encoding="utf-8", errors="ignore")
                diag["docx_raw_read"] = {"ok": bool(raw and raw.strip()), "len": len(raw) if raw else 0}
                if raw and raw.strip():
                    return {"text": raw, "source": "docx_converted_raw", "diagnostics": diag}
            except Exception as e:
                diag["docx_raw_read"] = {"ok": False, "len": 0, "note": str(e)}
        return {"text": "", "source": None, "diagnostics": diag}

    # 5) images -> Google Vision
    if suf in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        diag["vision_available"] = {"ok": VISION_AVAILABLE, "note": str(_vision_import_error) if not VISION_AVAILABLE else ""}
        if VISION_AVAILABLE:
            t = ocr_image_with_google_vision(filepath)
            diag["vision_image"] = {"ok": bool(t and t.strip()), "len": len(t) if t else 0}
            if t and t.strip():
                return {"text": t, "source": "vision_image", "diagnostics": diag}
            else:
                diag["vision_image_note"] = "vision returned empty text"
        else:
            diag["vision_image_note"] = "vision not available"
        return {"text": "", "source": None, "diagnostics": diag}

    # 6) fallback raw read
    try:
        raw = p.read_text(encoding="utf-8", errors="ignore")
        diag["raw_read"] = {"ok": bool(raw and raw.strip()), "len": len(raw) if raw else 0}
        if raw and raw.strip():
            return {"text": raw, "source": "raw_read", "diagnostics": diag}
    except Exception as e:
        diag["raw_read"] = {"ok": False, "len": 0, "note": str(e)}

    return {"text": "", "source": None, "diagnostics": diag}

# ---------- Semantic-aware chunking ----------
def _is_heading(line: str) -> bool:
    line = line.strip()
    if not line:
        return False
    if re.match(r'^(section|article|chapter|clause)\b', line, flags=re.I):
        return True
    if re.match(r'^\d+(\.\d+)*\s*[-:]?', line):
        return True
    if line.isupper() and 1 <= len(line.split()) <= 8:
        return True
    return False

def semantic_chunk_text(text: str, target_chars: int = 1000, overlap: int = 200) -> List[Tuple[str, str]]:
    if not text or not text.strip():
        return []
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    paras = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
    blocks: List[str] = []
    i = 0
    while i < len(paras):
        p = paras[i]
        first_line = p.split('\n', 1)[0].strip()
        if _is_heading(first_line) and (i + 1) < len(paras):
            combined = first_line + "\n\n" + paras[i + 1].strip()
            blocks.append(combined)
            i += 2
        else:
            blocks.append(p)
            i += 1
    if len(blocks) <= 2 and len(text) > target_chars * 1.5:
        chunks = []
        idx = 0
        n = len(text)
        while idx < n:
            end = min(idx + target_chars, n)
            chunk = text[idx:end].strip()
            if chunk:
                chunks.append((str(uuid.uuid4()), chunk))
            idx += target_chars - overlap
        return chunks
    merged: List[str] = []
    current = ""
    for b in blocks:
        if not current:
            current = b
        elif len(current) + len(b) + 2 <= target_chars:
            current = current + "\n\n" + b
        else:
            merged.append(current.strip())
            current = b
            if len(current) > target_chars * 1.5:
                subs = re.split(r'(?<=[\.\?\!])\s+', current)
                acc = ""
                for s in subs:
                    if len(acc) + len(s) + 1 <= target_chars:
                        acc = (acc + " " + s).strip() if acc else s
                    else:
                        if acc:
                            merged.append(acc.strip())
                        acc = s
                current = acc
    if current:
        merged.append(current.strip())
    chunks: List[str] = []
    for m in merged:
        if len(m) <= target_chars:
            chunks.append(m)
        else:
            idx = 0
            n = len(m)
            while idx < n:
                end = min(idx + target_chars, n)
                chunk = m[idx:end].strip()
                if chunk:
                    chunks.append(chunk)
                idx += target_chars - overlap
    out: List[Tuple[str, str]] = []
    seen = set()
    for c in chunks:
        key = c[:300]
        if key in seen:
            continue
        seen.add(key)
        out.append((str(uuid.uuid4()), c))
    return out

def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Tuple[str, str]]:
    return semantic_chunk_text(text, target_chars=chunk_size, overlap=overlap)

# ---------- DB helpers ----------
def create_user_if_not_exists(user_id: str, username: Optional[str] = None):
    if not user_id:
        return
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM users WHERE user_id = ?", (user_id,))
    if cur.fetchone():
        return
    cur.execute("INSERT INTO users (user_id, username, created_at) VALUES (?, ?, ?)",
                (user_id, username or None, datetime.utcnow().isoformat()))
    conn.commit()

def associate_thread_with_user(user_id: str, thread_id: str):
    if not thread_id:
        return
    if user_id:
        create_user_if_not_exists(user_id)
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM threads WHERE thread_id = ?", (thread_id,))
    if cur.fetchone():
        cur.execute("UPDATE threads SET user_id = ? WHERE thread_id = ?", (user_id, thread_id))
    else:
        cur.execute("INSERT INTO threads (thread_id, user_id, file_name, filepath, ingested_at) VALUES (?, ?, ?, ?, ?)",
                    (thread_id, user_id, None, None, None))
    conn.commit()

def get_threads_for_user(user_id: str):
    if not user_id:
        return []
    cur = conn.cursor()
    cur.execute("SELECT thread_id, file_name, filepath, ingested_at FROM threads WHERE user_id = ?", (user_id,))
    rows = cur.fetchall()
    return [{"thread_id": r[0], "file_name": r[1], "filepath": r[2], "ingested_at": r[3]} for r in rows]

def get_user_for_thread(thread_id: str):
    if not thread_id:
        return None
    cur = conn.cursor()
    cur.execute("SELECT user_id FROM threads WHERE thread_id = ?", (thread_id,))
    r = cur.fetchone()
    return r[0] if r else None

# ---------- Utility: legal / lexical analysis ----------
LEGAL_KEYWORDS = [
    "agreement", "party", "parties", "indemnify", "indemnity", "warranty", "liability",
    "governing law", "jurisdiction", "termination", "clause", "herein", "hereby",
    "force majeure", "arbitration", "confidentiality", "non-disclosure", "nda"
]

def _simple_clean_tokens(text: str):
    tokens = re.findall(r"\b[A-Za-z\-']+\b", text)
    return [t.lower() for t in tokens]

def detect_legal_like(text: str, threshold: int = 3) -> bool:
    low = text.lower()
    count = sum(1 for kw in LEGAL_KEYWORDS if kw in low)
    return count >= threshold

def detect_hard_words(text: str, top_n: int = 40) -> List[str]:
    tokens = _simple_clean_tokens(text)
    counts = Counter(tokens)
    scored = []
    for w, c in counts.items():
        if len(w) < 4:
            continue
        freq_score = None
        if zipf_frequency is not None:
            try:
                freq = zipf_frequency(w, "en")
                freq_score = freq
            except Exception:
                freq_score = None
        rarity = (0 if freq_score is None else (7.0 - freq_score))
        length_factor = max(0, (len(w) - 6) / 10.0)
        score = rarity * 1.2 + length_factor + math.log(c + 1)
        scored.append((w, score, c, freq_score))
    scored.sort(key=lambda x: (-x[1], -x[2]))
    terms = [t[0] for t in scored[:top_n]]
    stop_like = {"section", "shall", "including", "included", "thereof", "hereby"}
    terms = [t for t in terms if t not in stop_like]
    return terms

def _read_ingested_filepath(user_id: Optional[str], thread_id: str) -> Optional[str]:
    try:
        _, _, file_record = _chat_paths(user_id, thread_id)
        if file_record.exists():
            rec = json.loads(file_record.read_text(encoding='utf-8'))
            return rec.get("filepath")
    except Exception:
        pass
    return None

# ---------- LLM helpers ----------
def _call_model_system_then_user(system_prompt: str, user_prompt: str, temperature: Optional[float] = None) -> str:
    sys = SystemMessage(content=system_prompt)
    hum = HumanMessage(content=user_prompt)
    try:
        if temperature is None:
            resp = model.invoke([sys, hum])
            return getattr(resp, "content", str(resp))
        else:
            try:
                tmp_model = ChatGoogleGenerativeAI(
                    model=os.getenv('GOOGLE_MODEL', DEFAULT_GOOGLE_MODEL),
                    temperature=float(temperature),
                    google_api_key=os.getenv('GOOGLE_API_KEY')
                )
                resp = tmp_model.invoke([sys, hum])
                return getattr(resp, "content", str(resp))
            except Exception:
                resp = model.invoke([sys, hum])
                return getattr(resp, "content", str(resp))
    except Exception as e:
        return f"(model error: {e})"

def _call_model_with_messages(messages: List[BaseMessage], temperature: Optional[float] = None):
    try:
        if temperature is None:
            return model.invoke(messages)
        else:
            try:
                tmp_model = ChatGoogleGenerativeAI(
                    model=os.getenv('GOOGLE_MODEL', DEFAULT_GOOGLE_MODEL),
                    temperature=float(temperature),
                    google_api_key=os.getenv('GOOGLE_API_KEY')
                )
                return tmp_model.invoke(messages)
            except Exception:
                return model.invoke(messages)
    except Exception as e:
        return SystemMessage(content=f"(model error: {e})")

# ---------- Ingest / retrieve / index ----------
def ingest_file(filepath: str, file_name: str, thread_id: str, user_id: Optional[str] = None):
    """
    Ingest a file into the per-user-per-thread FAISS index.
    Returns success + diagnostics (on failure).
    """
    if user_id:
        create_user_if_not_exists(user_id)
        associate_thread_with_user(user_id, thread_id)

    index, metadata = _ensure_index_for_chat(user_id, thread_id)

    # Attempt extraction with diagnostics
    extraction = extract_text_with_diagnostics(filepath)
    text = extraction.get("text", "") or ""
    source = extraction.get("source")
    diagnostics = extraction.get("diagnostics", {})

    if not text or not text.strip():
        return {"success": False, "message": "No text extracted from file.", "diagnostics": diagnostics, "source": source}

    # proceed with chunking / embeddings
    chunks = _chunk_text(text, chunk_size=1000, overlap=200)
    texts = [c[1] for c in chunks]
    if not texts:
        return {"success": False, "message": "No chunks created from file.", "diagnostics": diagnostics, "source": source}

    vectors = _embed_model.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    start_idx = int(index.ntotal)
    index.add(vectors.astype(np.float32))
    for i, (chunk_id, chunk_text) in enumerate(chunks):
        idx_key = start_idx + i
        metadata[str(int(idx_key))] = {
            "file_name": file_name,
            "chunk_id": chunk_id,
            "text": chunk_text[:4000]
        }
    _persist_index_for_chat(user_id, thread_id, index, metadata)

    # write ingested file record
    _, _, file_record = _chat_paths(user_id, thread_id)
    with open(file_record, "w", encoding="utf-8") as f:
        json.dump({"file_name": file_name, "filepath": str(filepath), "user_id": user_id, "extracted_from": source}, f, ensure_ascii=False, indent=2)

    # update threads table
    cur = conn.cursor()
    now = datetime.utcnow().isoformat()
    cur.execute("SELECT 1 FROM threads WHERE thread_id = ?", (thread_id,))
    if cur.fetchone():
        cur.execute("UPDATE threads SET user_id = ?, file_name = ?, filepath = ?, ingested_at = ? WHERE thread_id = ?",
                    (user_id, file_name, str(filepath), now, thread_id))
    else:
        cur.execute("INSERT INTO threads (thread_id, user_id, file_name, filepath, ingested_at) VALUES (?, ?, ?, ?, ?)",
                    (thread_id, user_id, file_name, str(filepath), now))
    conn.commit()
    return {"success": True, "message": f"Ingested {len(chunks)} chunks (source={source}) into chat {thread_id} for user {user_id}", "diagnostics": diagnostics, "source": source}

def thread_has_ingested_file(thread_id: str, user_id: Optional[str] = None) -> bool:
    _, _, file_record = _chat_paths(user_id, thread_id)
    if file_record.exists():
        return True
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM threads WHERE thread_id = ? AND file_name IS NOT NULL", (thread_id,))
    return cur.fetchone() is not None

# ---------- Reranking / retrieval ----------
def rerank_candidates_with_cross_encoder(query: str, candidates: List[Dict[str, Any]], top_k: int = FINAL_TOP_K) -> List[Dict[str, Any]]:
    if not candidates:
        return []
    if _cross_encoder is None:
        return candidates[:top_k]
    pairs = [(query, c.get("text", "")) for c in candidates]
    try:
        scores = _cross_encoder.predict(pairs, convert_to_numpy=True)
    except Exception:
        return candidates[:top_k]
    for c, s in zip(candidates, scores):
        c["score"] = float(s)
    candidates.sort(key=lambda x: x.get("score", 0.0), reverse=True)
    return candidates[:top_k]

def retrieve_similar_chunks(query: str, user_id: Optional[str], thread_id: str, top_k: int = 3):
    index, metadata = _ensure_index_for_chat(user_id, thread_id)
    if index.ntotal == 0:
        return []
    q_vec = _embed_model.encode([query], convert_to_numpy=True)
    ann_k = min(ANN_TOP_K, int(index.ntotal))
    try:
        D, I = index.search(q_vec.astype(np.float32), ann_k)
    except Exception:
        return []
    candidates = []
    for idx in I[0]:
        if int(idx) < 0:
            continue
        meta = metadata.get(str(int(idx)), {})
        candidates.append({"index": int(idx), "file_name": meta.get("file_name"), "text": meta.get("text", "")})
    if not candidates:
        return []
    final = rerank_candidates_with_cross_encoder(query, candidates, top_k=top_k)
    out = []
    for c in final:
        out.append({"index": c.get("index"), "file_name": c.get("file_name"), "text": c.get("text"), **({"score": c.get("score")} if "score" in c else {})})
    return out

def retrieve_all_threads(user_id: Optional[str] = None):
    if not user_id:
        return []
    cur = conn.cursor()
    cur.execute("SELECT thread_id FROM threads WHERE user_id = ?", (user_id,))
    rows = cur.fetchall()
    return [r[0] for r in rows]

# ---------- Quick analyze & get_term_context ----------
def quick_analyze(filepath: str, user_id: Optional[str], thread_id: str) -> Dict[str, Any]:
    try:
        extraction = extract_text_with_diagnostics(filepath)
        text = extraction.get("text", "") or ""
        if not text or not text.strip():
            return {"success": False, "message": "No text extracted for analysis.", "diagnostics": extraction.get("diagnostics", {})}
        n = len(text)
        first = text[:2000]
        middle = text[max(0, n//2 - 1000): min(n, n//2 + 1000)]
        last = text[-2000:]
        sample = "\n\n---\n\n".join([first, middle, last])
        system_prompt = (
            "You are a concise legal document analyst.\n"
            "Using only the provided document excerpts, produce: (A) a short plain-English summary aimed at a non-lawyer, (B) three short factual bullets labelled FACTS that are explicitly supported by the excerpts, and (C) a single-line confidence indicator (High/Medium/Low)."
        )
        user_prompt = f"Document excerpts:\n{sample}\n\nReturn: (1) 3-sentence summary, (2) three FACTS bullets, (3) Confidence:"
        summary_text = _call_model_system_then_user(system_prompt, user_prompt, temperature=0.2)
        tokens = _simple_clean_tokens(text)
        counts = Counter(tokens)
        common = [w for w, _ in counts.most_common(30) if len(w) > 4][:15]
        legal_like = detect_legal_like(text)
        return {"success": True, "summary": summary_text, "keywords": common, "legal_like": legal_like, "sample_snippets": {"first": first[:2000], "middle": middle[:2000], "last": last[:2000]}, "diagnostics": extraction.get("diagnostics", {})}
    except Exception as e:
        return {"success": False, "message": f"quick_analyze error: {e}"}

def get_term_context(user_id: Optional[str], thread_id: str, term: str) -> Dict[str, Any]:
    try:
        filepath = _read_ingested_filepath(user_id, thread_id)
        if not filepath:
            return {"success": False, "message": "No ingested file for this thread."}
        extraction = extract_text_with_diagnostics(filepath)
        text = extraction.get("text", "") or ""
        if not text:
            return {"success": False, "message": "No text available."}
        m = re.search(r"([^.]{0,300}\b" + re.escape(term) + r"\b[^.]{0,300}\.)", text, flags=re.I | re.S)
        paragraph = m.group(0) if m else text[:400]
        q_prompt_sys = (
            "You are a legal assistant. Explain the requested term in plain English *in the context of the provided paragraph*. Keep the definition short (1-2 sentences), give one concrete example relevant to the document, and list any immediate legal implications or actions (one line). If the paragraph does not define the term, say 'Not stated in document.'"
        )
        q_user = f"Term: {term}\n\nParagraph:\n{paragraph}\n\nDefinition:"
        definition = _call_model_system_then_user(q_prompt_sys, q_user, temperature=0.0)
        return {"success": True, "term": term, "definition": definition, "snippet": paragraph}
    except Exception as e:
        return {"success": False, "message": f"get_term_context error: {e}"}

# ---------- Chat integration ----------
class ChatState(TypedDict):
    messages: Annotated[list[BaseMessage], add_messages]

def _extract_ids_and_text_from_human(msg_content: str):
    if not msg_content:
        return None, None, msg_content
    uid = None
    tid = None
    rest = msg_content
    if msg_content.startswith("[USER_ID:"):
        try:
            close = msg_content.index("]")
            uid = msg_content[9:close].strip()
            rest = msg_content[close+1:].lstrip()
        except ValueError:
            rest = msg_content
    if rest.startswith("[THREAD_ID:"):
        try:
            close2 = rest.index("]\n")
            tid = rest[11:close2].strip()
            rest = rest[close2+2:]
        except ValueError:
            if "]" in rest:
                close2 = rest.index("]")
                tid = rest[11:close2].strip()
                rest = rest[close2+1:].lstrip()
    return uid, tid, rest

def deep_analyze_short_variant(filepath, user_id, thread_id):
    return quick_analyze(filepath, user_id, thread_id)

def quick_analyze_short_variant(filepath, user_id, thread_id):
    return quick_analyze(filepath, user_id, thread_id)

def chat_node(state: ChatState):
    messages = state.get('messages', [])
    last_user_msg = None
    last_user_index = None
    for i in range(len(messages)-1, -1, -1):
        m = messages[i]
        if isinstance(m, HumanMessage):
            last_user_msg = m
            last_user_index = i
            break
    if last_user_msg is None:
        return {"messages": []}

    uid, tid, cleaned_text = _extract_ids_and_text_from_human(last_user_msg.content)
    if tid is not None and uid is None:
        err = SystemMessage(content=("Access denied: A [USER_ID:<id>] is required to continue this thread."))
        return {"messages": [err]}
    if tid is not None and uid is not None:
        owner = get_user_for_thread(tid)
        if owner is not None and owner != uid:
            err = SystemMessage(content=("Access denied: This thread does not belong to the specified USER_ID."))
            return {"messages": [err]}

    if tid is None:
        new_messages = messages
        try:
            resp = _call_model_with_messages(new_messages, temperature=0.0)
            return {"messages": [resp] if not isinstance(resp, list) else resp}
        except Exception as e:
            err_msg = SystemMessage(content=f"Error calling model: {e}")
            return {"messages": [err_msg]}
    else:
        cleaned_human = HumanMessage(content=cleaned_text)
        new_messages_list = messages.copy()
        new_messages_list[last_user_index] = cleaned_human
        retrieved = retrieve_similar_chunks(cleaned_text, user_id=uid, thread_id=tid, top_k=4)
        if retrieved:
            context_texts = []
            for r in retrieved:
                fn = r.get("file_name", "document")
                snippet = r.get("text", "") or ""
                snippet_short = snippet[:1200].replace('\n', ' ')
                context_texts.append(f"--- From file: {fn} ---\n{snippet_short}\n")
            context_combined = "\n\n".join(context_texts)
            system_prompt_text = _build_strict_system_prompt(context_combined, max_context_chars=5000)
            system_prompt = SystemMessage(content=system_prompt_text)
            new_messages = [system_prompt] + new_messages_list
            try:
                resp = _call_model_with_messages(new_messages, temperature=0.0)
                return {"messages": [resp] if not isinstance(resp, list) else resp}
            except Exception as e:
                err_msg = SystemMessage(content=f"Error calling model: {e}")
                return {"messages": [err_msg]}
        else:
            system_prompt = SystemMessage(content=FALLBACK_SYSTEM_PROMPT)
            new_messages = [system_prompt] + new_messages_list
            try:
                resp = _call_model_with_messages(new_messages, temperature=0.0)
                return {"messages": [resp] if not isinstance(resp, list) else resp}
            except Exception as e:
                err_msg = SystemMessage(content=f"Error calling model: {e}")
                return {"messages": [err_msg]}

# compile graph
graph = StateGraph(ChatState)
graph.add_node("chat_node", chat_node)
graph.add_edge(START, "chat_node")
graph.add_edge("chat_node", END)
chatbot = graph.compile(checkpointer=checkpointer)

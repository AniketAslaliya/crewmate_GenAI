
# backend_rag/form_processing.py
from __future__ import annotations
import json
import re
# Near the top of api_server.py
from typing import List, Optional, Dict, Any # Ensure List is included


from .models import model, call_model_system_then_user

# --- Required Imports ---
from docx import Document
from PyPDF2 import PdfReader, PdfWriter # Keep for placeholder/future use

# --- Assuming these are available from your existing setup ---
from .models import call_model_system_then_user
from .retrieval import retrieve_similar_chunks # For context
# Assuming OcrResult class is defined or you handle OCR output directly
# from .ocr import OcrResult # Example if defined elsewhere

# Placeholder Class Definitions (if not defined elsewhere)
from typing import List # Make sure List is imported
from pydantic import BaseModel 
class OcrWord(BaseModel):
    text: str
    bbox: List[int] # [xmin, ymin, xmax, ymax] coordinates

class OcrPage(BaseModel):
    page_number: int
    width: int
    height: int
    words: List[OcrWord]

class DetailedOcrResult(BaseModel):
    pages: List[OcrPage]

# --- Field Detection & Classification ---

# In backend_rag/form_processing.py

# --- (Keep existing imports: json, re, typing, LLM call, etc.) ---
# --- (Ensure DetailedOcrResult, OcrPage, OcrWord classes are defined/imported) ---
# --- (Ensure find_word_bbox and find_nearby_blank_bbox helpers are present) ---

def _generate_field_description(field_label: str, field_type: str) -> str:
    """Uses LLM to generate a short, helpful description for a form field."""
    print(f"--- [Form Description] Generating description for: {field_label} ({field_type}) ---")
    
    # Simple prompt for description
    system_prompt = (
        "You provide the description that must be meaningful and clear. It should be slightly long, providing context (1-2 sentences).\n"
        "Based on the field label and type, provide its description \n"
        "explaining what the user should enter.\n"
        "Example Input: Label='Date of Birth', Type='date'\n"
        "Example Output: Enter the date you were born, usually in DD/MM/YYYY format.\n"
        "Example Input: Label='PAN No.', Type='pan'\n"
        "Example Output: Enter your 10-character Permanent Account Number (PAN).\n"
        "Rules:\n"
        "- Be user-friendly.\n"
        "- Output ONLY the description string."
    )
    user_prompt = f"Field Label: \"{field_label}\"\nField Type: \"{field_type}\"\n\nDescription:"
    
    try:
        description = call_model_system_then_user(system_prompt, user_prompt, temperature=0.1)
        return description.strip()
    except Exception as e:
        print(f"--- [Form Description] Failed: {e} ---")
        return "Enter the required information for this field." # Fallback

# In backend_rag/form_processing.py

import pdfplumber





# In backend_rag/form_processing.py
# In backend_rag/form_processing.py

# In backend_rag/form_processing.py

def _batch_generate_descriptions(fields_info: List[Dict], context: str) -> Dict[str, str]:
    """
    Generates DETAILED, SIMPLE explanations for form fields.
    """
    if not fields_info:
        return {}
        
    print(f"--- [Form Batch Desc] Generating simple, detailed descriptions... ---")
    
    field_list_str = "\n".join([f"- Label: \"{f['label_text']}\", Type: \"{f['semantic_type']}\"" for f in fields_info])
    safe_context = context[:4000] if context else "No specific document context provided."

    # --- UPDATED PROMPT FOR SIMPLICITY ---
    system_prompt = (
        "You are a helpful AI assistant explaining a legal form to a layperson. "
        "Your goal is to explain **what to write** and **why it is needed** in extremely simple, plain language.\n\n"
        
        "**INSTRUCTIONS:**\n"
        "1. **Simplify:** Use 5th-grade reading level English. No legal jargon.\n"
        "2. **Explain:** For each field, explain what information is required and *why* the form needs it (based on the document context).\n"
        "3. **Tone:** Be helpful and encouraging.\n"
        "4. Return a valid JSON object: `{\"Label Text\": \"Description\"}`.\n\n"
        
        "**Examples:**\n"
        "- Instead of 'Provide DOB for identification', say: 'Please enter your Date of Birth. This helps us confirm exactly who you are.'\n"
        "- Instead of 'Enter Statute Limitation Date', say: 'Enter the date the incident happened. This is important to make sure the form is filed on time.'\n\n"
        
        "**Document Context:**\n"
        f"{safe_context}\n\n"
        
        "**Field List:**\n"
        f"{field_list_str}\n\n"
        
        "**JSON Output:**"
    )
    
    user_prompt = "Generate simple JSON descriptions."
    
    descriptions = {}
    try:
        response = call_model_system_then_user(
            system_prompt, user_prompt, temperature=0.3, model_instance=model
        )
        match = re.search(r'\{.*\}', response, re.DOTALL)
        if match:
            descriptions = json.loads(match.group(0))
            print(f"--- [Form Batch Desc] Generated {len(descriptions)} simple descriptions. ---")
        else:
            print("--- [Form Batch Desc] Failed: LLM did not return valid JSON object.")
    except Exception as e:
        print(f"--- [Form Batch Desc] Failed: {e} ---")
        descriptions = {f['label_text']: "Please write the required information here." for f in fields_info}
        
    return descriptions


def detect_form_fields(ocr_result: DetailedOcrResult, context_summary: str = "") -> List[Dict]:
    """
    V3: Detects fields via LLM using layout data AND generates context-aware descriptions.
    """
    print("--- [Form Processing V3] Detecting fields with context... ---")
    
    # 1. Serialize OCR data (Same as before)
    page_data_for_llm = []
    word_count = 0
    for page in ocr_result.pages:
        if word_count > 4000: break
        words_list = [w.model_dump() for w in page.words]
        page_data_for_llm.append({
            "page_number": page.page_number,
            "width": page.width,
            "height": page.height,
            "words": words_list
        })
        word_count += len(page.words)
    
    if word_count == 0: return []

    # 2. LLM Call to Identify Fields (Same structure as before)
    system_prompt_detect = (
        "You are an expert form layout analyst. Identify all fillable fields in this form JSON.\n"
        "For each field, find the `label_text` and the `bbox` of the BLANK input area.\n"
        "Return a valid JSON list `[{'label_text':..., 'semantic_type':..., 'bbox':..., 'page_number':...}]`."
    )
    llm_input_json = json.dumps(page_data_for_llm, default=int)[:100000]
    user_prompt_detect = f"Form OCR Data:\n---\n{llm_input_json}\n---\n\nIdentified Fields JSON List:"

    detected_fields_final = []
    try:
        response = call_model_system_then_user(
            system_prompt_detect, user_prompt_detect, temperature=0.0, model_instance=model
        )
        
        match = re.search(r'\[.*\]', response, re.DOTALL)
        if not match: return []
        
        llm_fields = json.loads(match.group(0))

        # 3. Prepare list for batch description
        temp_field_list_for_desc = []
        parsed_fields_temp = {} 
        
        for i, field_data in enumerate(llm_fields):
            if not isinstance(field_data.get('bbox'), list) or len(field_data['bbox']) != 4: continue
            if 'page_number' not in field_data: continue

            field_id = f"field_{i}"
            label_text = field_data.get('label_text', 'Unknown')
            semantic_type = field_data.get('semantic_type', 'text').lower()
            
            temp_field_list_for_desc.append({"label_text": label_text, "semantic_type": semantic_type})
            parsed_fields_temp[field_id] = field_data 

        # 4. Single LLM Call to Generate CONTEXT-AWARE Descriptions
        # Pass the context_summary here!
        field_descriptions = _batch_generate_descriptions(temp_field_list_for_desc, context_summary)

        # 5. Compile Final Field List
        for field_id, field_data in parsed_fields_temp.items():
            label_text = field_data['label_text']
            semantic_type = field_data['semantic_type']
            page_number = field_data['page_number']
            
            is_sensitive = semantic_type in ['pan', 'address', 'phone', 'email', 'amount']
            
            # Use the generated detailed description
            description = field_descriptions.get(label_text, "Enter the required information.") 

            field = {
                'id': field_id,
                'label_text': label_text,
                'bbox': field_data['bbox'], 
                'page_number': page_number,
                'semantic_type': semantic_type,
                'confidence': 'High', 
                'is_sensitive': is_sensitive,
                'description': description,
                'value': ""
            }
            detected_fields_final.append(field)

    except Exception as e:
        print(f"--- [Form Processing V3] Error: {e} ---")

    return detected_fields_final
        

# --- Suggestion Generation ---

def generate_field_suggestions(field: Dict, context: str) -> List[str]:
    """
    Generates potential values for a specific form field using Gemini/LLM.
    """
    field_type = field.get('semantic_type', 'other')
    field_label = field.get('label_text', '')
    is_sensitive = field.get('is_sensitive', False)

    if is_sensitive or field_type in ['signature', 'pan', 'other']:
         print(f"--- [Form Suggestion] Skipping suggestions for sensitive/unsuitable field: {field_label} ({field_type}) ---")
         return []

    print(f"--- [Form Suggestion] Generating suggestions for: {field_label} ({field_type}) ---")
    safe_context = context # TODO: Implement proper redaction for sensitive context

    system_prompt = (
        "You are a concise assistant that suggests potential values for a form field.\n"
        "Rules:\n"
        "- Based on the field type, label, and context, provide 1-3 likely candidate values.\n"
        "- Return a valid JSON list of strings `[]`.\n"
        "- Keep suggestions short and appropriate for the field type.\n"
        "- Output ONLY the JSON list."
    )
    user_prompt = (
        f"Field Label: \"{field_label}\"\n"
        f"Field Type: \"{field_type}\"\n"
        f"Context Summary: \"{safe_context[:1000]}\"\n\n"
        f"Task: Provide 1-3 candidate values as a JSON list:"
    )

    suggestions = []
    try:
        response = call_model_system_then_user(system_prompt, user_prompt, temperature=0.4)
        match = re.search(r'\[.*\]', response, re.DOTALL)
        if match:
            raw_suggestions = json.loads(match.group(0))
            suggestions = [str(s) for s in raw_suggestions if isinstance(s, (str, int, float))][:3]
            print(f"--- [Form Suggestion] Generated suggestions: {suggestions} ---")
    except Exception as e:
        print(f"--- [Form Suggestion] Error generating suggestions: {e} ---")

    return suggestions

# --- Form Export ---

# --- Option 1: Filling PDFs (Placeholder - Requires PDF with AcroForm fields) ---
def fill_pdf_form(template_path: str, output_path: str, field_data: Dict[str, str]):
    """Placeholder: Fills a PDF form template with interactive fields."""
    print(f"Placeholder: Filling PDF {template_path} -> {output_path}. Needs implementation if PDF AcroForms are used.")
    # Basic implementation using PyPDF2 (commented out):
    # try:
    #     reader = PdfReader(template_path)
    #     writer = PdfWriter()
    #     for page in reader.pages:
    #         writer.add_page(page)
    #     form_fields = reader.get_fields()
    #     if form_fields:
    #         writer.update_page_form_field_values(writer.pages[0], field_data)
    #         with open(output_path, "wb") as output_stream:
    #             writer.write(output_stream)
    #         print(f"--- [Form Export] Successfully created filled PDF: {output_path} ---")
    #     else:
    #         print(f"--- [Form Export] WARNING: No interactive fields found in PDF {template_path}.")
    #         # Copy original if no fields
    #         with open(template_path, "rb") as f_in, open(output_path, "wb") as f_out: f_out.write(f_in.read())
    # except Exception as e:
    #     print(f"--- [Form Export] ERROR filling PDF: {e} ---")
    #     raise
    pass # Keep as pass if not implemented

# --- Option 2: Filling DOCX (Functional Implementation) ---
def fill_docx_form(template_path: str, output_path: str, field_data: Dict[str, str]):
    """
    Fills a DOCX template by replacing placeholder text like {{field_id}}.
    """
    print(f"--- [Form Export] Filling DOCX {template_path} -> {output_path} ---")
    try:
        document = Document(template_path)
        
        # Replace placeholders in paragraphs
        for para in document.paragraphs:
            for key, value in field_data.items():
                placeholder = f"{{{{{key}}}}}" # e.g., {{complainant_name}}
                if placeholder in para.text:
                    # Simple replacement logic (might need refinement for complex cases)
                    inline = para.runs
                    for run in inline:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        # Replace placeholders in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in field_data.items():
                             placeholder = f"{{{{{key}}}}}"
                             if placeholder in para.text:
                                 inline = para.runs
                                 for run in inline:
                                     if placeholder in run.text:
                                         run.text = run.text.replace(placeholder, str(value))

        document.save(output_path)
        print(f"--- [Form Export] Successfully created filled DOCX: {output_path} ---")

    except Exception as e:
        print(f"--- [Form Export] ERROR filling DOCX: {e} ---")
        raise # Re-raise the exception to be caught by the API